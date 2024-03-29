# Import the required libraries
from tensorflow.keras.models import load_model
from tensorflow.keras.preprocessing import image
import numpy as np
import streamlit as st
import google.generativeai as genai
import os
from PIL import Image
import time
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX

# Confiure the Generative AI API using genai
genai.configure(api_key = "AIzaSyAIewGMqAtMEtZMZjDJgEPNEwh_Q74yfGw")

# Setting the prompt for the GenerativeAI Model
prompt = "You are an expert in Opthamology(Expert in treating with the eye disease) You are provided with the stages of the disease Diabetic Retinopathy from Stage 0 to 4 at the end of the line, You should then respond with a set of precautionary measures, recommended medications, and healthy habits tailored to the specific stage of diabetic retinopathy indicated by the user. Stage Input is : (stage)" 

# Function to get the gemini response
def get_gemini_response(stage):
    prompt_input = prompt.replace("(stage)", str(stage))
    # Call the Generative AI Model
    genai_model = genai.GenerativeModel('gemini-pro')
    response = genai_model.generate_content(prompt_input)
    return response.text

# Function to load and preprocess the image
def load_and_preprocess_image(img_input, target_size=(512, 512)):
    img = image.load_img(img_input, target_size=target_size)
    img_tensor = image.img_to_array(img)
    img_tensor = np.expand_dims(img_tensor, axis=0)
    img_tensor /= 255.0
    return img_tensor

# Summary generator used for display as word wise
def summary_generator():
    for sentence in result.split(" "):
        yield sentence + " "
        time.sleep(0.10)

# Function to convert the generated response into a Document
def save_as_docx(text, filename):
    doc = Document()
    doc.add_paragraph(text)
    key_point = doc.add_paragraph("This Generated Medication report is based on our model and it's previous processing. Make sure that you consult your doctor for further medication 🙂")
    paragraph_format = key_point.paragraph_format
    paragraph_format.fill = WD_COLOR_INDEX.YELLOW
    paragraph_format.text = WD_COLOR_INDEX.BLACK
    doc.save(filename)

# Load the trained model
model = load_model('model/diabetic_retinopathy_model_version_check.h5')

# Initialize the streamlit App
st.set_page_config(page_title = "DR Detection")
st.title("Detect your DR Stage.")
st.header("Diabetic Retinopathy Detection using Deep Learning")
image_input = st.file_uploader("Choose your image...", type = ['jpg', 'jpeg', 'png'])
image_txt = ""
if image_input is not None:
    image_txt = Image.open(image_input) # Reading the image
    st.image(image_txt, caption = "Uploaded Image", use_column_width = True) # Display the uploaded image
predict_button = st.button("Predict the stage of Diabetic Retinopathy")

if predict_button:
    test_image = load_and_preprocess_image(image_input)
    prediction = model.predict(test_image)
    stage = np.argmax(prediction)
    st.subheader(f'The Predicted Intensity is {stage}')
    result = get_gemini_response(stage)
    st.subheader("Generated Medication is :")
    st.write_stream(summary_generator)
    st.caption("This Generated Medication report is based on our model and it's previous processing. Make sure that you consult your doctor for further medication 🙂")
    save_as_docx(result, "DR_REPORT.docx")
    st.download_button(label = "Download Your report",
                       data = open("DR_REPORT.docx", "rb"),
                       file_name = "DR_REPORT.docx",
                       mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
